import copy
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from build_application_package import build_resume


FIXTURE = Path(__file__).parents[2] / "packages" / "test-fixtures" / "src" / "application-package.json"


class ApplicationDocumentBuilderTests(unittest.TestCase):
    def setUp(self):
        self.payload = json.loads(FIXTURE.read_text(encoding="utf-8"))

    def test_builds_a_letter_document_with_real_headings_and_claims(self):
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "resume.docx"
            build_resume(self.payload, output)
            document = Document(output)
            self.assertEqual(round(document.sections[0].page_width.inches, 1), 8.5)
            self.assertEqual(round(document.sections[0].left_margin.inches, 1), 1.0)
            headings = [paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Heading 1"]
            self.assertEqual(headings, ["Profile", "Selected experience", "Technical focus"])
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("reduced fixture triage time by 20%", text)

    def test_rejects_an_unsupported_metric(self):
        payload = copy.deepcopy(self.payload)
        payload["sections"][0]["items"][0]["text"] = "Reduced fixture triage time by 99%."
        with tempfile.TemporaryDirectory() as directory:
            with self.assertRaisesRegex(ValueError, "unsupported metric"):
                build_resume(payload, Path(directory) / "resume.docx")

    def test_rejects_an_unknown_evidence_reference(self):
        payload = copy.deepcopy(self.payload)
        payload["sections"][0]["items"][0]["evidenceIds"] = [999]
        with tempfile.TemporaryDirectory() as directory:
            with self.assertRaisesRegex(ValueError, "verified evidence"):
                build_resume(payload, Path(directory) / "resume.docx")


if __name__ == "__main__":
    unittest.main()
