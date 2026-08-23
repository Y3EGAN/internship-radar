"""Build evidence-bound DOCX application artifacts from a validated JSON manifest."""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x55, 0x55, 0x55)


def set_font(run, size: float, color=INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_id = max([int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))] or [0]) + 1
    num_id = max([int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))] or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "bullet"), ("w:lvlText", "•"), ("w:lvlJc", "left")):
        element = OxmlElement(tag)
        element.set(qn("w:val"), value)
        level.append(element)
    ppr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.extend((spacing, tabs, indent))
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, number))
    ppr.append(num_pr)
    set_font(paragraph.add_run(text), 11)


def validate_manifest(payload: dict) -> None:
    evidence = {int(item["id"]): item["fact"] for item in payload.get("evidence", [])}
    if not evidence:
        raise ValueError("verified evidence is required")
    for section in payload.get("sections", []):
        for claim in section.get("items", []):
            ids = claim.get("evidenceIds", [])
            if not ids or any(int(item) not in evidence for item in ids):
                raise ValueError("every resume claim must reference verified evidence")
            facts = " ".join(evidence[int(item)] for item in ids).casefold()
            for metric in re.findall(r"\b\d+(?:\.\d+)?%?\b", claim.get("text", "")):
                if metric.casefold() not in facts:
                    raise ValueError("resume claim contains an unsupported metric")


def build_resume(payload: dict, output: Path) -> None:
    validate_manifest(payload)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    configure_styles(document)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    candidate = payload["candidate"]
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    set_font(title.add_run(candidate["name"]), 24, INK, True)
    contact = document.add_paragraph()
    contact.paragraph_format.space_after = Pt(14)
    set_font(contact.add_run(" | ".join(candidate.get("contactLines", []))), 10, MUTED)
    summary_heading = document.add_paragraph("Profile", style="Heading 1")
    summary_heading.paragraph_format.keep_with_next = True
    document.add_paragraph(payload["summary"])
    num_id = add_bullet_numbering(document)
    for content in payload.get("sections", []):
        document.add_paragraph(content["heading"], style="Heading 1")
        for claim in content.get("items", []):
            add_bullet(document, claim["text"], num_id)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Evidence-bound application package"), 8, MUTED, italic=True)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("manifest", type=Path)
    parser.add_argument("--resume-out", type=Path, required=True)
    args = parser.parse_args()
    payload = json.loads(args.manifest.read_text(encoding="utf-8"))
    build_resume(payload, args.resume_out)


if __name__ == "__main__":
    main()
